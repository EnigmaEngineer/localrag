"""Document parsers for different file formats."""

from pathlib import Path

from langchain_core.documents import Document
from loguru import logger


def parse_file(file_path: Path) -> list[Document]:
    """Parse a file into a list of LangChain Documents.

    Each document represents a logical unit (e.g., a page in a PDF).
    """
    suffix = file_path.suffix.lower()
    parser = PARSERS.get(suffix)

    if parser is None:
        raise ValueError(f"No parser available for {suffix}")

    return parser(file_path)


def _parse_pdf(file_path: Path) -> list[Document]:
    """Parse PDF using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    documents = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "page": i + 1,
                        "total_pages": len(reader.pages),
                        "file_type": "pdf",
                    },
                )
            )

    logger.debug(f"PDF parsed: {file_path.name} → {len(documents)} pages")
    return documents


def _parse_docx(file_path: Path) -> list[Document]:
    """Parse DOCX using python-docx."""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())

    return [
        Document(
            page_content=full_text,
            metadata={
                "source": file_path.name,
                "file_type": "docx",
            },
        )
    ]


def _parse_text(file_path: Path) -> list[Document]:
    """Parse plain text or markdown files."""
    text = file_path.read_text(encoding="utf-8")

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path.name,
                "file_type": file_path.suffix.lstrip("."),
            },
        )
    ]


def _parse_csv(file_path: Path) -> list[Document]:
    """Parse CSV files — each row becomes a document."""
    import csv

    documents = []
    with open(file_path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for i, row in enumerate(reader):
            text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "row": i + 1,
                        "file_type": "csv",
                    },
                )
            )

    return documents


PARSERS = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".txt": _parse_text,
    ".md": _parse_text,
    ".csv": _parse_csv,
}
